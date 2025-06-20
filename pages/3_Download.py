import streamlit as st
import pdfplumber
import re
from docx import Document
import tempfile
import uuid

st.set_page_config(page_title="Step 3: Download", layout="centered")
st.title("📥 Step 3: Download Enhanced Resume")

# JD Templates
jd_templates = {
    "qa engineer": "Selenium, JIRA, REST API, Postman, Agile, TestNG, JUnit",
    "python developer": "Python, Django, Flask, MySQL, REST API, Git, Docker",
    "business analyst": "SQL, Requirements Gathering, Stakeholders, JIRA, Agile",
    "frontend developer": "HTML, CSS, JavaScript, React, Figma, Git",
    "devops engineer": "AWS, Docker, Jenkins, CI/CD, Terraform, Linux"
}

def extract_text_from_pdf(file):
    with pdfplumber.open(file) as pdf:
        return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])

def clean(text):
    return re.sub(r"[^a-zA-Z0-9 ]", " ", text).lower()

def extract_keywords(text):
    return set(clean(text).split())

def enhance_resume(resume_text, jd_text):
    jd_keywords = extract_keywords(jd_text)
    resume_keywords = extract_keywords(resume_text)
    missing = jd_keywords - resume_keywords
    doc = Document()
    doc.add_heading(st.session_state.name, 0)
    doc.add_paragraph(f"Email: {st.session_state.email} | Phone: +91-9876543210 | Location: India")

    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph("Experienced professional with relevant experience. Proficient in " + ", ".join(list(missing)[:6]) + ".")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(sorted(jd_keywords)))

    doc.add_heading("Experience", level=1)
    doc.add_paragraph("Software Engineer, ABC Corp (2019 - Present)\n- Designed and executed test plans\n- Collaborated with cross-functional teams")

    doc.add_heading("Education", level=1)
    doc.add_paragraph("B.Tech in Computer Science\nXYZ University, 2015 - 2019")
    return doc

if "payment_done" not in st.session_state:
    st.warning("⚠️ Please complete Step 2 first.")
else:
    resume_text = extract_text_from_pdf(st.session_state.resume_file)
    jd_input = st.session_state.role_or_jd.strip().lower()
    auto_jd = jd_templates.get(jd_input, st.session_state.role_or_jd)
    enhanced_doc = enhance_resume(resume_text, auto_jd)
    temp_path = os.path.join(tempfile.gettempdir(), f"enhanced_{uuid.uuid4()}.docx")
    enhanced_doc.save(temp_path)

    with open(temp_path, "rb") as f:
        st.download_button("📄 Download Enhanced Resume (DOCX)", f, file_name="Enhanced_Resume.docx")
